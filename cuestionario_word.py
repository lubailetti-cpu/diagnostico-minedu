"""
Genera el cuestionario completo en formato Word (.docx).
Muestra todas las preguntas con la bifurcación entre MINEDU sede central, DRE y UGEL.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Estructura: (seccion, pregunta, tipo, opciones, aplica_a, condicional)
PREGUNTAS = [
    # === SECCIÓN A: PERFIL ===
    ("A. Perfil del servidor",
     "¿En qué nivel del sector educativo trabaja?",
     "Select desplegable",
     ["MINEDU sede central",
      "Dirección Regional de Educación (DRE)",
      "Unidad de Gestión Educativa Local (UGEL)"],
     "Todos",
     "Pregunta FILTRO. Define el camino que sigue el resto del cuestionario."),

    ("A. Perfil del servidor",
     "¿A qué régimen laboral pertenece?",
     "Select desplegable",
     ["D. L. N° 1057 - CAS (Contrato Administrativo de Servicios)",
      "D.L. N° 276 (Carrera Administrativa)",
      "D. L. N° 728 (Régimen de la actividad privada)",
      "Ley N° 30057 (Ley del Servicio Civil)",
      "Locador de servicios",
      "Practicante"],
     "Todos", ""),

    ("A. Perfil del servidor",
     "¿En qué categoría se ubica su unidad organizacional?",
     "Select desplegable",
     ["Alta Dirección",
      "Órganos de Control y Defensa",
      "Órganos de Asesoramiento",
      "Órganos de Apoyo",
      "Órganos de Línea — Gestión Pedagógica",
      "Órganos de Línea — Gestión Institucional",
      "Organismos Adscritos"],
     "MINEDU",
     "Solo si nivel = MINEDU. Basado en ROF DS 001-2015-MINEDU."),

    ("A. Perfil del servidor",
     "¿Cuál unidad específicamente?",
     "Select desplegable (cascada)",
     ["(Lista filtrada según la categoría elegida arriba, entre 1 y 26 opciones)"],
     "MINEDU",
     "Solo si nivel = MINEDU. DEBEDSAR Sede Central y COAR siguen separadas."),

    ("A. Perfil del servidor",
     "¿A qué Dirección Regional de Educación (DRE) pertenece?",
     "Select desplegable",
     ["DRE Amazonas", "DRE Áncash", "DRE Apurímac", "DRE Arequipa", "DRE Ayacucho",
      "DRE Cajamarca", "DRE Callao", "DRE Cusco", "DRE Huancavelica", "DRE Huánuco",
      "DRE Ica", "DRE Junín", "DRE La Libertad", "DRE Lambayeque",
      "DRE Lima Metropolitana", "DRE Lima Provincias", "DRE Loreto",
      "DRE Madre de Dios", "DRE Moquegua", "DRE Pasco", "DRE Piura",
      "DRE Puno", "DRE San Martín", "DRE Tacna", "DRE Tumbes", "DRE Ucayali"],
     "DRE", "Solo si nivel = DRE."),

    ("A. Perfil del servidor",
     "¿A qué Unidad de Gestión Educativa Local (UGEL) pertenece?",
     "Campo de texto libre",
     ["(Escribir el nombre de su UGEL)"],
     "UGEL", "Solo si nivel = UGEL."),

    ("A. Perfil del servidor",
     "¿Cuál es el rol principal que ejerce actualmente dentro de su U.O.?",
     "Select desplegable",
     ["Asesor/a", "Coordinador/a", "Especialista técnico/a", "Especialista administrativo/a", "Otras"],
     "Todos", ""),

    ("A. Perfil del servidor",
     "En general, ¿cuántos años de experiencia tiene en el sector público?",
     "Select desplegable",
     ["Menos de 2 años", "De 2 a 5 años", "De 6 a 10 años", "Más de 10 años"],
     "Todos", ""),

    # === SECCIÓN B: TU TRABAJO DOCUMENTAL ===
    ("B. Tu trabajo documental",
     "Distribución de su tiempo laboral semanal",
     "Barra segmentada interactiva (3 puntos arrastrables, 4 colores). La suma debe ser 100%.",
     ["1. Funciones sustantivas de tu unidad o equipo (análisis técnico, propuestas de mejora, asesoría especializada, diseño de proyectos, coordinación estratégica)",
      "2. Elaboración o revisión de documentos normativos e instrumentos de gestión (directivas, normas técnicas, lineamientos, ROF, MOF, MAPRO, fichas, diagramas BPMN)",
      "3. Tareas administrativas repetitivas (informes de conformidad, oficios, memorandos, ayudas memoria, reportes, atención de correos)",
      "4. Atención de demandas externas no planificadas (solicitudes de Congreso, Contraloría, PCM, GOREs, CCVs)"],
     "Todos", "Si no mueve los puntos, aparece popup que pregunta si realmente las 4 actividades ocupan el mismo tiempo (25% cada una)."),

    ("B. Tu trabajo documental",
     "Para cada tipo de documento, indique cuántos elabora o revisa al mes y cuánto tiempo le toma en promedio.",
     "Matriz por tipo (cards verticales con pills). Para cada uno de los 9 tipos, dos preguntas: Cantidad y Tiempo.",
     ["TIPOS:",
      "  - Ficha de proceso o de producto",
      "  - Ficha de procedimiento",
      "  - Diagrama de procesos (BPMN)",
      "  - Informe técnico",
      "  - Presentación institucional",
      "  - Ayuda memoria / resumen",
      "  - Oficio, memorando, conformidad",
      "  - Documento normativo (directiva, lineamiento)",
      "  - Guía metodológica",
      "",
      "OPCIONES CANTIDAD AL MES: 0 / 1-2 / 3-5 / 6-10 / 11+",
      "OPCIONES TIEMPO POR CADA UNO: <30 min / 30 min-2 h / 2-4 h / 4-8 h / 1-3 días / 3+ días"],
     "Todos", "NÚCLEO DE LA LÍNEA DE BASE."),

    ("B. Tu trabajo documental",
     "¿Cuál es la principal dificultad al elaborar estos instrumentos?",
     "Checkboxes (máximo 2 opciones)",
     ["La normativa está dispersa o es de difícil acceso",
      "No existen plantillas o modelos actualizados",
      "El proceso requiere muchas revisiones y aprobaciones",
      "Falta de tiempo suficiente",
      "Falta de capacitación o conocimiento técnico",
      "Las instrucciones o criterios no son claros",
      "Recibo documentos derivados incorrectamente",
      "Otras"],
     "Todos", ""),

    ("B. Tu trabajo documental",
     "¿Tiene acceso a plantillas, guías o herramientas estandarizadas para elaborar estos instrumentos?",
     "Opción única",
     ["Sí, y las uso regularmente",
      "Sí, pero están desactualizadas o son insuficientes",
      "No tengo acceso a herramientas de este tipo",
      "No las necesito"],
     "Todos", ""),

    # === SECCIÓN C: REPROCESO Y CALIDAD ===
    ("C. Reproceso y calidad",
     "De cada 10 documentos que entrega, ¿cuántos le devuelven con observaciones?",
     "Opción única",
     ["Casi ninguno (0-1)",
      "Pocos (2-3)",
      "La mitad (4-6)",
      "Bastantes (7-8)",
      "Casi todos (9-10)"],
     "Todos", ""),

    ("C. Reproceso y calidad",
     "Cuando un documento requiere reproceso, ¿cuántas rondas de revisión/devolución se producen en promedio antes de su aprobación final?",
     "Opción única",
     ["1 ronda", "2 rondas", "3 rondas", "Más de 3 rondas"],
     "Todos", ""),

    ("C. Reproceso y calidad",
     "Cuando le devuelven un documento, ¿por qué suele ser?",
     "Checkboxes (máximo 2 opciones)",
     ["Errores de forma (formato, redacción, ortografía, citas)",
      "Errores de fondo (criterio técnico, análisis incompleto)",
      "Falta de alineamiento normativo",
      "Estilo institucional inadecuado",
      "Información incompleta o desactualizada",
      "Otras"],
     "Todos", ""),

    # === SECCIÓN D: BÚSQUEDA DE INFORMACIÓN ===
    ("D. Búsqueda de información",
     "En promedio, ¿cuánto tiempo dedica a buscar normativa, guías metodológicas o modelos de referencia ANTES de empezar a elaborar un instrumento?",
     "Opción única",
     ["Menos de 30 minutos",
      "Entre 30 minutos y 1 hora",
      "Entre 1 y 2 horas",
      "Más de 2 horas",
      "No suelo buscarlo, trabajo de memoria o con lo que tengo"],
     "Todos", ""),

    ("D. Búsqueda de información",
     "¿Qué tan trabajoso le resulta asegurarse de que la normativa, guía o modelo que está usando es la versión vigente?",
     "Escala 1-4 con colores semáforo",
     ["1 (verde) = Muy fácil",
      "2",
      "3",
      "4 (rojo) = Muy trabajoso"],
     "Todos", ""),

    # === SECCIÓN E: RITMO Y DEMANDAS EXTERNAS ===
    ("E. Ritmo y demandas externas",
     "¿Con qué frecuencia recibe solicitudes urgentes de entidades externas (Congreso, Contraloría, PCM, Defensoría, GOREs u otras) que interrumpen su planificación semanal?",
     "Opción única",
     ["Nunca", "1 a 2 veces por semana", "3 a 5 veces por semana", "Diariamente o más"],
     "Todos", ""),

    ("E. Ritmo y demandas externas",
     "Cuando recibe estas solicitudes externas, ¿cuál es el plazo habitual que le otorgan para responder?",
     "Opción única",
     ["Más de 3 días", "Entre 1 y 3 días", "Menos de 24 horas",
      "Menos de 4 horas (media jornada)", "No aplica / no recibo solicitudes externas"],
     "Todos", ""),

    # === SECCIÓN F: OPORTUNIDADES EN TU ÁREA ===
    ("F. Oportunidades en tu área",
     "Valore qué tan identificado se siente con cada afirmación. Escala 1 (totalmente en desacuerdo, verde) a 4 (totalmente de acuerdo, rojo).",
     "8 afirmaciones, cada una con escala 1-4 colores semáforo",
     ["1. La carga documental y administrativa ha ocasionado retrasos en el cumplimiento de actividades planificadas en el POI de mi unidad",
      "2. La carga de trabajo de mi equipo NO está distribuida de la mejor manera",
      "3. Me siento agotado/a al terminar mi semana laboral",
      "4. Me siento poco satisfecho/a con la calidad de los entregables que terminamos produciendo como área",
      "5. Hay mucho trabajo manual repetitivo en mi área que podría automatizarse",
      "6. Mucha información importante de mi área está en la cabeza de pocas personas",
      "7. Tenemos datos importantes pero dispersos en distintos formatos que nos cuesta consolidar",
      "8. Nos cuesta transformar la información disponible en indicadores para tomar decisiones"],
     "Todos", ""),

    ("F. Oportunidades en tu área",
     "En el último mes, ¿cuántas veces ha tenido que extender su jornada laboral para cumplir con entregables documentales o administrativos?",
     "Opción única",
     ["Nunca", "1 a 2 veces", "3 a 5 veces", "Más de 5 veces"],
     "Todos", ""),

    ("F. Oportunidades en tu área",
     "Si pudiera recuperar tiempo de las tareas repetitivas y administrativas, ¿a qué actividades lo destinaría?",
     "Checkboxes (máximo 2 opciones)",
     ["Análisis y diagnóstico de problemas",
      "Mejora de procesos e innovación",
      "Coordinación con otras unidades de organización o equipos",
      "Atención a ciudadanos o usuarios del MINEDU",
      "Capacitación y desarrollo profesional",
      "Otras"],
     "Todos", ""),

    # === SECCIÓN G: USO ACTUAL DE IA ===
    ("G. Uso actual de inteligencia artificial",
     "¿Cuál es tu nivel actual de uso de inteligencia artificial?",
     "Opción única",
     ["No uso ninguna herramienta de IA",
      "Nivel 1 - Básico: consultas y consejos puntuales",
      "Nivel 2 - Pragmático: resumir, traducir, redactar, buscar información",
      "Nivel 3 - Constructor: automatizo tareas, armo asistentes, trabajo con archivos",
      "Nivel 4 - Avanzado: agentes o flujos integrados a mis sistemas"],
     "Todos", ""),

    ("G. Uso actual de inteligencia artificial",
     "¿Usas alguna herramienta de inteligencia artificial para facilitar tu trabajo?",
     "Opción única",
     ["Sí, con cuenta pagada por el MINEDU",
      "Sí, con cuenta pagada por mí",
      "Sí, con cuenta gratuita",
      "No uso ninguna"],
     "Todos", ""),

    ("G. Uso actual de inteligencia artificial",
     "¿Para qué tareas de tu trabajo usas inteligencia artificial hoy?",
     "Checkboxes (máximo 3 opciones)",
     ["Redactar borradores de informes u oficios",
      "Resumir documentos largos",
      "Buscar información o normativa",
      "Traducir textos",
      "Generar ideas o lluvia inicial",
      "Analizar datos en Excel",
      "Hacer presentaciones",
      "Otra",
      "No uso IA para mi trabajo"],
     "Todos", "Solo aparece si en la pregunta anterior respondió que SÍ usa IA."),

    # === SECCIÓN H: PERCEPCIÓN PLATAFORMA ===
    ("H. Percepción sobre una plataforma de asistencia con IA",
     "Valore su nivel de acuerdo con las siguientes afirmaciones sobre una posible plataforma de asistencia con IA para elaborar instrumentos de gestión.",
     "4 afirmaciones, cada una con escala 1-5 (1 = totalmente en desacuerdo, 5 = totalmente de acuerdo)",
     ["1. Reduciría significativamente el tiempo que dedico a tareas repetitivas",
      "2. Me ayudaría a generar borradores de mayor calidad que los que elaboro sin apoyo",
      "3. Me daría más tiempo para actividades estratégicas y de mayor valor",
      "4. Me gustaría participar en una prueba piloto de esta herramienta"],
     "Todos", ""),

    ("H. Percepción sobre una plataforma de asistencia con IA",
     "¿Cuáles serían sus principales preocupaciones al usar una herramienta de IA para elaborar instrumentos de gestión?",
     "Checkboxes (máximo 2 opciones)",
     ["Errores técnicos o normativos en los borradores",
      "Pérdida del criterio profesional especializado",
      "No se adapta a las particularidades del MINEDU",
      "Problemas de seguridad o confidencialidad de la información",
      "Resistencia institucional o falta de validación oficial",
      "Tiempo de capacitación necesario",
      "No tengo preocupaciones, estoy abierto/a a probarla",
      "Otras"],
     "Todos", ""),

    ("H. Percepción sobre una plataforma de asistencia con IA",
     "¿Qué condición considera indispensable para adoptar una herramienta de IA en su trabajo?",
     "Checkboxes (máximo 2 opciones)",
     ["Que esté validada y autorizada institucionalmente",
      "Que haya una capacitación previa",
      "Que sea fácil de usar (interfaz intuitiva)",
      "Que garantice la confidencialidad de la información",
      "Que exista soporte técnico disponible",
      "Que permita revisar y modificar el borrador antes de utilizarlo oficialmente"],
     "Todos", ""),

    # === SECCIÓN I: DRE / UGEL ===
    ("I. Trabajo específico DRE / UGEL",
     "¿Cuántas Instituciones Educativas (IIEE) están a tu cargo o jurisdicción?",
     "Opción única",
     ["Menos de 10", "10 a 50", "51 a 200", "Más de 200"],
     "DRE/UGEL", ""),

    ("I. Trabajo específico DRE / UGEL",
     "¿Cuántas visitas o supervisiones a IIEE realizas en promedio al mes?",
     "Opción única",
     ["Ninguna", "1 a 3", "4 a 10", "Más de 10"],
     "DRE/UGEL", ""),

    ("I. Trabajo específico DRE / UGEL",
     "Después de una visita a IIEE, ¿cuánto tiempo te toma elaborar el informe correspondiente?",
     "Opción única",
     ["Menos de 1 hora", "Entre 1 y 3 horas", "Entre 3 y 8 horas", "Más de un día"],
     "DRE/UGEL", ""),

    ("I. Trabajo específico DRE / UGEL",
     "Antes de una visita, ¿cuánto tiempo dedicas a preparar antecedentes y documentación de la IIEE?",
     "Opción única",
     ["Menos de 30 minutos", "Entre 30 minutos y 2 horas", "Más de 2 horas", "No preparo antecedentes"],
     "DRE/UGEL", ""),

    ("I. Trabajo específico DRE / UGEL",
     "¿Cuánto tiempo dedicas a consolidar respuestas que vienen de IIEE en formatos distintos (Excel, Word, PDF)?",
     "Opción única",
     ["Menos de 2 horas por semana", "Entre 2 y 5 horas por semana",
      "Entre 5 y 10 horas por semana", "Más de 10 horas por semana"],
     "DRE/UGEL", ""),

    ("I. Trabajo específico DRE / UGEL",
     "¿Cuántas normas o directivas del MINEDU traduces o adaptas para las IIEE al mes?",
     "Opción única",
     ["Ninguna", "1 a 3", "4 a 10", "Más de 10"],
     "DRE/UGEL", ""),

    ("I. Trabajo específico DRE / UGEL",
     "¿Cuánto tiempo dedicas a responder consultas frecuentes de directores y docentes?",
     "Opción única",
     ["Menos de 2 horas por semana", "Entre 2 y 5 horas por semana",
      "Entre 5 y 10 horas por semana", "Más de 10 horas por semana"],
     "DRE/UGEL", ""),

    ("I. Trabajo específico DRE / UGEL",
     "¿Para cuáles tareas de tu unidad sería MÁS útil una herramienta con IA?",
     "Checkboxes (máximo 3 opciones)",
     ["Preparar visitas de supervisión",
      "Elaborar informes post-visita",
      "Consolidar reportes que vienen de IIEE",
      "Traducir normas a lenguaje claro para directores",
      "Responder consultas frecuentes de docentes",
      "Generar oficios masivos a IIEE",
      "Procesar trámites de docentes (contratos, evaluaciones)",
      "Capacitaciones y materiales pedagógicos",
      "Otra"],
     "DRE/UGEL", ""),

    # === SECCIÓN J: TU VOZ CUENTA ===
    ("J. Tu voz cuenta",
     "En sus propias palabras: ¿cuáles son las 2 o 3 tareas repetitivas que más tiempo le quitan en su trabajo semanal?",
     "Texto libre",
     ["(Pregunta abierta)"],
     "Todos", ""),
]


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_shaded_paragraph(doc, text, color_hex="EFF4FF"):
    """Párrafo con sombreado de fondo (para el badge de 'Aplica a')."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    p_pr.append(shd)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(9)
    return p


def color_aplica(aplica):
    if aplica == "Todos":
        return "F0FDF4"
    elif aplica == "MINEDU":
        return "DBEAFE"
    elif aplica == "DRE":
        return "FEF3C7"
    elif aplica == "UGEL":
        return "FCE7F3"
    elif aplica == "DRE/UGEL":
        return "FEF3C7"
    return "FFFFFF"


def label_aplica(aplica):
    return {
        "Todos": "APLICA A: TODOS",
        "MINEDU": "APLICA A: SOLO MINEDU SEDE CENTRAL",
        "DRE": "APLICA A: SOLO DRE",
        "UGEL": "APLICA A: SOLO UGEL",
        "DRE/UGEL": "APLICA A: SOLO DRE / UGEL"
    }.get(aplica, aplica)


def main():
    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== Portada =====
    title = doc.add_heading('Diagnóstico de carga operativa MINEDU', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    subtitle = doc.add_paragraph('Cuestionario completo — versión con bifurcación MINEDU / DRE / UGEL')
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.italic = True
    subtitle.runs[0].font.color.rgb = RGBColor(0x5C, 0x61, 0x78)

    doc.add_paragraph()

    intro = doc.add_paragraph()
    intro.add_run("Iniciativa de innovación propuesta por la ").font.size = Pt(11)
    run = intro.add_run("Unidad de Organización y Métodos (UNOME)")
    run.bold = True
    run.font.size = Pt(11)
    intro.add_run(".").font.size = Pt(11)

    confid = doc.add_paragraph()
    r = confid.add_run("Esta encuesta es confidencial. No estamos recogiendo datos personales. Los resultados se comunicarán únicamente de forma agregada, sin identificar respuestas individuales.")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x16, 0x55, 0x34)

    doc.add_paragraph()

    # ===== Leyenda =====
    add_heading(doc, "Leyenda de colores por aplicabilidad", level=2,
                color=RGBColor(0x1E, 0x3A, 0x5F))

    legend_table = doc.add_table(rows=5, cols=2)
    legend_table.style = 'Light Grid Accent 1'
    legend_data = [
        ("APLICA A: TODOS", "Pregunta que ve cualquier servidor (MINEDU, DRE o UGEL)"),
        ("APLICA A: SOLO MINEDU SEDE CENTRAL", "Solo si en la pregunta 1 selecciona MINEDU"),
        ("APLICA A: SOLO DRE", "Solo si en la pregunta 1 selecciona DRE"),
        ("APLICA A: SOLO UGEL", "Solo si en la pregunta 1 selecciona UGEL"),
        ("APLICA A: SOLO DRE / UGEL", "Aparece tanto para DRE como para UGEL (bloque específico de supervisión a IIEE)"),
    ]
    for i, (label, desc) in enumerate(legend_data):
        cell0 = legend_table.cell(i, 0)
        cell0.text = label
        for p in cell0.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
        legend_table.cell(i, 1).text = desc
        for p in legend_table.cell(i, 1).paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

    doc.add_paragraph()

    # ===== Resumen de cantidades por perfil =====
    add_heading(doc, "Cuántas preguntas ve cada perfil", level=2,
                color=RGBColor(0x1E, 0x3A, 0x5F))

    total_minedu = sum(1 for p in PREGUNTAS if p[4] in ("Todos", "MINEDU"))
    total_dre = sum(1 for p in PREGUNTAS if p[4] in ("Todos", "DRE", "DRE/UGEL"))
    total_ugel = sum(1 for p in PREGUNTAS if p[4] in ("Todos", "UGEL", "DRE/UGEL"))

    summary = doc.add_table(rows=4, cols=3)
    summary.style = 'Light Grid Accent 1'
    summary.cell(0, 0).text = "Perfil"
    summary.cell(0, 1).text = "Preguntas que ve"
    summary.cell(0, 2).text = "Tiempo estimado"
    summary.cell(1, 0).text = "MINEDU sede central"
    summary.cell(1, 1).text = str(total_minedu)
    summary.cell(1, 2).text = "~7-8 minutos"
    summary.cell(2, 0).text = "DRE (Dirección Regional)"
    summary.cell(2, 1).text = str(total_dre)
    summary.cell(2, 2).text = "~9-10 minutos"
    summary.cell(3, 0).text = "UGEL"
    summary.cell(3, 1).text = str(total_ugel)
    summary.cell(3, 2).text = "~9-10 minutos"

    for row in summary.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    for cell in summary.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

    doc.add_page_break()

    # ===== Preguntas =====
    last_section = ""
    n = 0
    for sec, preg, tipo, opciones, aplica, condicional in PREGUNTAS:
        if sec != last_section:
            add_heading(doc, sec, level=1, color=RGBColor(0x1E, 0x3A, 0x5F))
            last_section = sec
        n += 1

        # Badge "Aplica a"
        add_shaded_paragraph(doc, "  " + label_aplica(aplica) + "  ", color_aplica(aplica))

        # Pregunta
        p = doc.add_paragraph()
        run_n = p.add_run(f"{n}. ")
        run_n.bold = True
        run_n.font.size = Pt(12)
        run_q = p.add_run(preg)
        run_q.bold = True
        run_q.font.size = Pt(12)

        # Tipo
        tipo_p = doc.add_paragraph()
        tipo_run = tipo_p.add_run("Tipo de respuesta: ")
        tipo_run.italic = True
        tipo_run.font.size = Pt(9)
        tipo_run.font.color.rgb = RGBColor(0x5C, 0x61, 0x78)
        tipo_val = tipo_p.add_run(tipo)
        tipo_val.font.size = Pt(9)
        tipo_val.font.color.rgb = RGBColor(0x5C, 0x61, 0x78)

        # Opciones
        for op in opciones:
            op_p = doc.add_paragraph(style='List Bullet')
            r = op_p.add_run(op)
            r.font.size = Pt(10)

        # Condicional / nota
        if condicional:
            nota = doc.add_paragraph()
            nota_label = nota.add_run("Nota: ")
            nota_label.italic = True
            nota_label.font.size = Pt(9)
            nota_label.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
            nota_val = nota.add_run(condicional)
            nota_val.italic = True
            nota_val.font.size = Pt(9)
            nota_val.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

        doc.add_paragraph()  # espacio

    output = "Cuestionario_Diagnostico_MINEDU.docx"
    doc.save(output)
    print(f"Word generado: {output}")
    print(f"Total preguntas en el cuestionario maestro: {len(PREGUNTAS)}")
    print(f"  MINEDU sede central ve: {total_minedu}")
    print(f"  DRE ve: {total_dre}")
    print(f"  UGEL ve: {total_ugel}")


if __name__ == "__main__":
    main()
